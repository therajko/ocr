import os
import json
import docx
import re
from openai import OpenAI
import argparse
from typing import Dict, Any, Tuple
from docx import Document
from pathlib import Path

def read_word_template(file_path: str) -> str:
    """
    Read content from a Word document.
    
    Args:
        file_path: Path to the Word document
        
    Returns:
        String containing the document content
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Template file not found: {file_path}")
        
    doc = docx.Document(file_path)
    content = []
    
    for para in doc.paragraphs:
        content.append(para.text)
    
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            content.append(' | '.join(cells))
    
    return '\n'.join(content)

def generate_json_and_markdown(template_content: str, output_dir: str, model_name: str) -> Tuple[str, str]:
    """
    Generate a question JSON file and a markdown template from a document content.
    
    Args:
        template_content: Content of the template document
        output_dir: Directory to save the output files
        
    Returns:
        Tuple of paths to the created question JSON and markdown template files
    """
    
    # Create output directory if it doesn't exist
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    # Initialize OpenAI client
    if model_name=='openai':
        client = OpenAI(
            api_key=os.environ.get("OPENAI_API_KEY")
        )
        model = 'gpt-4o'
    else:
        client = OpenAI(
            base_url=os.environ.get("HF_LLAMA_GGUF_MODEL_ENDPOINT"),
            api_key=os.environ.get("HUGGINGFACE_KEY")
        )
        model = 'tgi'
    # Prompt to generate question set JSON
    question_prompt = """
    You are given a document template. Analyze it and create a structured JSON question set that captures 
    all the fields that need to be filled in. The JSON structure should follow this format:
    {
  "product_description": {
    "product_name": "Extract the complete product name including pharmacopeia standard",
    "strength": "Extract strength value with unit",
    "manufacturing_license": "Extract the complete license number",
    "shelf_life": "Extract shelf life with period",
    "batch_sizes": "Extract all batch sizes mentioned",
    "pack_styles": "Extract all packaging configurations",
    "storage_conditions": "Extract complete storage instructions"
  },
  "batches_manufactured": {
    "total_count": "Extract the total number of batches manufactured (numeric value)"
  },
  "manufacturing_formula": {
    "ingredients": [
      {
        "serial_number": "Extract the serial number in the formula table",
        "material_name": "Extract the complete name of the material/ingredient including standards",
        "vendor": "Extract the vendor name for the material",
        "quantity_small_batch": "Extract quantity for smaller batch size with units",
        "quantity_large_batch": "Extract quantity for larger batch size with units",
        "quantity_per_tablet": "Extract quantity per tablet with units"
      }
    ]
  },
  ### And so on for all sections in the document
}
    
    Identify all sections and fields in the document that require data extraction.
    
    Here's the template content:
    {template_content}
    
    Return only the JSON structure with no additional text or explanations.
    """
    
    # Prompt to generate markdown template
    markdown_prompt = f"""
    You are given a document template. Convert it to a clean, well-formatted markdown template.
    Any fields that need to be filled in should be enclosed in curly braces, like {{field_name}}.
    Preserve the structure, tables, headings, and formatting as much as possible.
    
    Here's the template content:
    {template_content}
    
    Return only the markdown template with no additional text or explanations.
    """
    
    # Generate JSON question set
    print("Generating JSON question set...")
    question_response = client.responses.create(
        model=model,
        messages=[{"role": "user", "content": question_prompt}],
        temperature=0.2
    )
    question_content = question_response.output_text
    
    # Extract JSON part (in case the model includes additional text)
    json_match = re.search(r'```json\s*(.*?)\s*```', question_content, re.DOTALL)
    if json_match:
        question_content = json_match.group(1)
    else:
        json_match = re.search(r'({.*})', question_content, re.DOTALL)
        if json_match:
            question_content = json_match.group(1)
    
    # Validate JSON
    question_json = json.loads(question_content)
    question_content = json.dumps(question_json, indent=2)
    
    # Save JSON
    json_path = os.path.join(output_dir, "transcription_metadata.json")
    with open(json_path, "w") as f:
        f.write(question_content)
    print(f"JSON question set saved to {json_path}")
    
    # Generate markdown template
    print("Generating markdown template...")
    markdown_response = client.responses.create(
        model=model,
        messages=[{"role": "user", "content": markdown_prompt}],
        temperature=0.2
    )
    markdown_content = markdown_response.output_text
    
    # Extract markdown part (in case the model includes additional text)
    md_match = re.search(r'```markdown\s*(.*?)\s*```', markdown_content, re.DOTALL)
    if md_match:
        markdown_content = md_match.group(1)
    
    # Save markdown
    md_path = os.path.join(output_dir, "pqr_template.md")
    with open(md_path, "w") as f:
        f.write(markdown_content)
    print(f"Markdown template saved to {md_path}")
    
    return json_path, md_path

def read_document_content(docx_path):
    """
    Read and standardize document content while preserving exact ordering of elements.
    
    Args:
        docx_path: Path to the Word document
        
    Returns:
        Tuple of (document_content, elements) where:
        - document_content: String containing the standardized document content
        - elements: List of tuples containing (element_type, content, position)
    """
    doc = Document(docx_path)
    content = []
    current_position = 0
    elements = []
    
    # Add paragraphs with their position
    for para in doc.paragraphs:
        if para.text.strip():  # Only add non-empty paragraphs
            elements.append(('paragraph', para.text, current_position))
            current_position += 1
    
    # Add tables with their position
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            table_content.append(' | '.join(cells))
        elements.append(('table', table_content, current_position))
        current_position += 1
    
    # Sort elements by their position to maintain original order
    elements.sort(key=lambda x: x[2])
    
    # Convert elements to content while preserving order
    for element_type, content_item, _ in elements:
        if element_type == 'paragraph':
            content.append(content_item)
        elif element_type == 'table':
            content.extend(content_item)
    
    document_content = '\n'.join(content)
    return document_content, elements

def get_model_response(client, model_name, prompt, temperature=0.2):
    """
    Get response from the model using the appropriate API call based on the model type.
    
    Args:
        client: OpenAI client instance
        model_name: Name of the model to use
        prompt: The prompt to send to the model
        temperature: Temperature parameter for the model (default: 0.2)
        
    Returns:
        The model's response text
    """
    
    if model_name == 'openai':
        response = client.responses.create(
            model='gpt-4o',
            input=prompt,
            temperature=temperature
        )
        return response.output_text
    else:
        response = client.chat.completions.create(
            model='tgi',
            messages=prompt,
            temperature=temperature
        )
        return response.choices[0].message.content

def extract_questions_from_docx(docx_path, model_name):
    """Extract questions from a Word document by analyzing its content."""
    # Read the document content
    document_content, _ = read_document_content(docx_path)
    
    # Initialize OpenAI client
    if model_name == 'openai':
        client = OpenAI(
            api_key=os.environ.get("OPENAI_API_KEY")
        )
    else:
        client = OpenAI(
            base_url=os.environ.get("HF_LLAMA_GGUF_MODEL_ENDPOINT"),
            api_key=os.environ.get("HUGGINGFACE_KEY")
        )
    
    # Prompt to generate questions from document content
    prompt = """You are given a document template for a Product Quality Review (PQR) report. 
Generate a structured set of questions that will systematically capture all required information to complete this report.

Analyze the template carefully to:
1. Identify all data fields, both explicit (with placeholders) and implicit (required by context)
2. Recognize dynamic sections that may require variable numbers of entries
3. Understand relationships between different sections of the document

For each identified data need:
- Create specific, direct questions to elicit the required information
- Include contextual guidance where appropriate
- Ensure questions account for dynamic data structures (arrays/lists of items)
- Use field names that align with a logical data structure

Output Format:
Return a comprehensive JSON structure organized by document sections:
{
  "product_information": [
    {
      "question": "What is the name of the product?",
      "field_name": "product_name",
      "data_type": "string",
      "is_array": false,
      "context": "The official product name as it appears in regulatory documentation"
    }
    // Additional questions for each section

Here's the document content:
""" + document_content
    
    # Generate questions using the model
    print("Generating questions from document content...")
    response_text = get_model_response(client, model_name, prompt)
    
    # Clean and parse the response
    response_text = re.search(r'\{.*\}', response_text, re.DOTALL).group(0)

    questions_json = json.loads(response_text)
    return questions_json

def create_markdown_template(docx_path, output_path, model_name):
    """Create a markdown template from the document content using OpenAI API."""
    # Read the document content
    document_content, elements = read_document_content(docx_path)
    
    # Initialize OpenAI client
    if model_name == 'openai':
        client = OpenAI(
            api_key=os.environ.get("OPENAI_API_KEY")
        )
    else:
        client = OpenAI(
            base_url=os.environ.get("HF_LLAMA_GGUF_MODEL_ENDPOINT"),
            api_key=os.environ.get("HUGGINGFACE_KEY")
        )
    
    # Prompt to generate markdown template
    prompt = """You are given a document template for a Product Quality Review (PQR) report. 
                 Convert this document into a completely flexible, data-driven markdown template that dynamically adjusts to any data structure provided.

                 Rules for conversion:
                 1. Create a template that makes NO assumptions about specific values or quantities (no hardcoded batch sizes, number of materials, etc.)
                 2. All columns, rows, and data points must be dynamically generated from the input data
                 3. Convert tables to markdown with column headers that are generic and can be populated from data
                 4. For repeating elements, use clear looping indicators with generic array notation
                 5. Use variable placeholder formats like {{product_name}} for single items and array notation for lists
                 6. Avoid any fixed numbers or values in column headers - these should be populated from data too
                 7. Make all tables completely flexible to handle any number of items and any values
                 8. Include comments explaining the data structure needed for each section

                 Critical requirements:
                 - Do NOT hardcode any batch sizes like "200,000" or "800,000" - these should be dynamic values
                 - Quantity columns should be generic and based on the provided batch sizes, not fixed numbers
                 - Each batch size should have its own quantity column that's generated dynamically
                 - Product specifications, materials, quantities, etc. should ALL be variable

    """ + document_content + """Return the markdown template with appropriate comments explaining dynamic sections."""
    
    # Generate markdown template using the model
    print("Generating markdown template...")
    response_text = get_model_response(client, model_name, prompt)
    
    # Clean up the response and save to file
    markdown_content = response_text.replace("```markdown", "").replace("```", "")
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(markdown_content)
    
    print(f"Markdown template saved to {output_path}")

def main(docx_path, output_dir, model_name):
    """Main function to process the Word template."""
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Define output file paths
    json_path = os.path.join(output_dir, "transcription_metadata.json")
    md_path = os.path.join(output_dir, "pqr_template.md")
    
    # Check if both output files already exist
    if os.path.exists(json_path) and os.path.exists(md_path):
        print("Template files already exist. Skipping processing.")
        return json_path, md_path
    
    # Extract questions from the Word document
    questions = extract_questions_from_docx(docx_path, model_name)
    
    if not questions:
        print("Error: No questions were generated from the document.")
        return None, None
    
    # Create JSON output
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump({"questions": questions}, f, indent=2)
    
    # Create markdown template using OpenAI API
    create_markdown_template(docx_path, md_path, model_name)
    
    print(f"Template processing complete. Files saved to:")
    print(f"- JSON: {json_path}")
    print(f"- Markdown: {md_path}")
    
    return json_path, md_path

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description='Process Word template and generate JSON and markdown files')
    parser.add_argument('--docx_path', type=str, required=True, help='Path to the Word template file')
    parser.add_argument('--output_dir', type=str, required=True, help='Directory to save output files')
    parser.add_argument('--model_name', type=str, required=False, default='local_llm', help='Model name')
    args = parser.parse_args()
    
    main(args.docx_path, args.output_dir, args.model_name) 